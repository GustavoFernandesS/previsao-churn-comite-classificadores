from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
ASSETS = REPORTS / "assets"
ASSETS.mkdir(exist_ok=True)

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 42, 55)
GRAY = RGBColor(90, 98, 108)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, italic=False, color=DARK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_font(run, size=9, color=GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, RGBColor(31, 77, 120), 8, 4),
    }
    for style_name, (size, color, before, after) in settings.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.text = "Projeto A3 - Inteligencia Artificial | "
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_font(run, size=9, color=GRAY)
    add_page_number(footer)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(115)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COMITE DE CLASSIFICADORES")
    set_font(run, size=26, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    run = p.add_run("Predicao de churn em clientes de telecomunicacoes")
    set_font(run, size=15, color=GRAY)

    for label, value in [
        ("Universidade", "Universidade Sao Judas Tadeu"),
        ("Disciplina", "Inteligencia Artificial"),
        ("Integrantes", "[Preencher nomes dos integrantes]"),
        ("Professor(a)", "[Preencher nome do(a) professor(a)]"),
        ("Turma", "[Preencher turma]"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        set_font(r, size=11, bold=True)
        r = p.add_run(value)
        set_font(r, size=11, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(95)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sao Paulo\n2026")
    set_font(r, size=11, color=GRAY)
    doc.add_page_break()


def add_paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead) :])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        set_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        hdr.cells[i].width = Inches(widths[i])
        hdr.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        set_cell_margins(hdr.cells[i])
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=9.5, bold=True, color=BLUE)

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, size=9.5, bold=(i == 0))
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=9, italic=True, color=GRAY)


def create_figures():
    data_path = ROOT / "data" / "WA_Fn-UseC_-Telco-Customer-Churn.csv"
    df = pd.read_csv(data_path)
    churn = df["Churn"].value_counts().reindex(["No", "Yes"])

    plt.figure(figsize=(7.2, 4.2))
    bars = plt.bar(["Nao cancelou", "Cancelou"], churn.values, color=["#2E74B5", "#D99036"])
    plt.title("Distribuicao da variavel alvo Churn", weight="bold")
    plt.ylabel("Quantidade de clientes")
    plt.grid(axis="y", alpha=0.2)
    for bar, value in zip(bars, churn.values):
        plt.text(bar.get_x() + bar.get_width() / 2, value + 80, f"{value:,}".replace(",", "."), ha="center")
    plt.tight_layout()
    churn_path = ASSETS / "distribuicao_churn.png"
    plt.savefig(churn_path, dpi=180, bbox_inches="tight")
    plt.close()

    results = pd.DataFrame(
        {
            "Modelo": ["Comite", "Naive Bayes", "KNN", "Arvore de Decisao", "SVM"],
            "Acuracia": [0.7743, 0.6948, 0.7779, 0.7970, 0.7935],
            "Precisao": [0.5603, 0.4589, 0.5822, 0.6679, 0.6627],
            "Recall": [0.6952, 0.8369, 0.5775, 0.4679, 0.4519],
            "F1-score": [0.6205, 0.5928, 0.5799, 0.5503, 0.5374],
        }
    ).set_index("Modelo")
    ax = results.plot(kind="bar", figsize=(9, 4.8), color=["#2E74B5", "#5B9BD5", "#D99036", "#70AD47"])
    ax.set_title("Comparacao das metricas por modelo", weight="bold")
    ax.set_xlabel("")
    ax.set_ylabel("Resultado")
    ax.set_ylim(0, 1)
    ax.grid(axis="y", alpha=0.2)
    ax.legend(ncol=4, loc="upper center", bbox_to_anchor=(0.5, -0.16), frameon=False)
    plt.xticks(rotation=15, ha="right")
    plt.tight_layout()
    metrics_path = ASSETS / "comparacao_metricas.png"
    plt.savefig(metrics_path, dpi=180, bbox_inches="tight")
    plt.close()
    return churn_path, metrics_path


def build_report():
    churn_path, metrics_path = create_figures()
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Resumo", level=1)
    add_paragraph(
        doc,
        "Este projeto desenvolveu um pipeline de aprendizado supervisionado para prever o cancelamento "
        "de clientes de uma empresa de telecomunicacoes. Foram treinados quatro classificadores "
        "individuais - KNN, Naive Bayes, SVM e Arvore de Decisao - e um comite por votacao suave. "
        "O comite apresentou o melhor F1-score (0,6205), superando o melhor modelo individual "
        "(Naive Bayes, com 0,5928), e demonstrou melhor equilibrio entre precisao e recall.",
    )
    add_paragraph(doc, "Palavras-chave: classificacao; ensemble; churn; aprendizado supervisionado; inteligencia artificial.", bold_lead="Palavras-chave:")

    doc.add_heading("1. Introducao", level=1)
    add_paragraph(
        doc,
        "A perda de clientes, conhecida como churn, representa um problema relevante para empresas "
        "de servicos recorrentes. Identificar antecipadamente clientes com maior risco de cancelamento "
        "permite direcionar campanhas de retencao, melhorar o atendimento e reduzir perdas de receita.",
    )
    add_paragraph(
        doc,
        "O objetivo deste trabalho e aplicar tecnicas de aprendizado supervisionado na construcao de "
        "um comite de classificadores. O desempenho do comite e comparado ao desempenho de modelos "
        "individuais, utilizando metricas adequadas para uma base com classes desbalanceadas.",
    )

    doc.add_heading("2. Descricao da base de dados", level=1)
    add_paragraph(
        doc,
        "Foi utilizada a base publica Telco Customer Churn, disponibilizada na plataforma Kaggle. "
        "Cada linha representa um cliente e cada coluna descreve caracteristicas demograficas, "
        "servicos contratados, cobrancas e situacao de cancelamento.",
    )
    add_table(
        doc,
        ["Caracteristica", "Descricao"],
        [
            ("Registros", "7.043 clientes"),
            ("Colunas", "21 colunas no arquivo original"),
            ("Target", "Churn: Yes para cancelamento e No para permanencia"),
            ("Features numericas", "SeniorCitizen, tenure, MonthlyCharges e TotalCharges"),
            ("Features categoricas", "Genero, contrato, servicos, pagamento e outras"),
        ],
        [1.8, 4.7],
    )
    doc.add_picture(str(churn_path), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figura 1 - Distribuicao da variavel alvo. Fonte: elaboracao propria.")
    add_paragraph(
        doc,
        "A classe positiva corresponde a 1.869 clientes (26,54%), enquanto 5.174 clientes (73,46%) "
        "nao cancelaram. Esse desbalanceamento torna insuficiente avaliar os modelos apenas pela acuracia.",
    )

    doc.add_heading("3. Definicao do problema", level=1)
    add_paragraph(
        doc,
        "O problema foi definido como uma classificacao binaria. A variavel alvo e Churn, convertida "
        "para 1 quando o cliente cancelou e 0 quando permaneceu. As variaveis independentes incluem "
        "tempo como cliente, cobranca mensal, cobranca total, tipo de contrato, forma de pagamento e "
        "servicos contratados.",
    )
    add_paragraph(
        doc,
        "Em termos de inferencia, o modelo busca identificar quais clientes apresentam maior risco de "
        "cancelamento com base em seus atributos. A classe positiva e o churn, pois representa o evento "
        "de maior interesse para uma estrategia de retencao.",
    )

    doc.add_heading("4. Pre-processamento", level=1)
    add_paragraph(doc, "O pre-processamento foi implementado por meio de pipelines para evitar vazamento de dados. As principais etapas foram:")
    add_bullets(
        doc,
        [
            "Conversao de TotalCharges de texto para valor numerico, revelando 11 valores ausentes.",
            "Imputacao dos valores numericos ausentes pela mediana.",
            "Remocao de customerID, pois a coluna identifica clientes, mas nao explica churn.",
            "Codificacao one-hot das variaveis categoricas.",
            "Padronizacao das variaveis numericas com StandardScaler.",
            "Separacao estratificada em 80% para treino e 20% para teste.",
        ],
    )
    add_paragraph(
        doc,
        "A divisao resultou em 5.634 registros de treino e 1.409 de teste. Depois da codificacao, "
        "cada conjunto passou a possuir 45 atributos e nenhum valor ausente.",
    )

    doc.add_heading("5. Modelagem dos classificadores individuais", level=1)
    add_paragraph(
        doc,
        "Foram selecionados quatro algoritmos com diferentes formas de aprendizado. O uso de modelos "
        "diversos e importante para a construcao do comite, pois cada algoritmo pode reconhecer "
        "padroes distintos nos dados.",
    )
    add_table(
        doc,
        ["Modelo", "Configuracao e finalidade"],
        [
            ("KNN", "Classificacao pelos 11 vizinhos mais proximos."),
            ("Naive Bayes", "Modelo probabilistico gaussiano, simples e sensivel aos sinais de churn."),
            ("SVM", "Kernel RBF com calibracao de probabilidades."),
            ("Arvore de Decisao", "Profundidade maxima 6 e minimo de 20 amostras por folha."),
        ],
        [1.8, 4.7],
    )
    add_paragraph(
        doc,
        "Todos os modelos foram treinados no mesmo conjunto e avaliados no mesmo conjunto de teste, "
        "mantendo uma comparacao justa. O pre-processamento foi ajustado somente com dados de treino.",
    )

    doc.add_heading("6. Construcao do comite", level=1)
    add_paragraph(
        doc,
        "O comite foi construido por votacao suave (soft voting), combinando KNN, Naive Bayes, SVM e "
        "Arvore de Decisao. Nessa estrategia, cada classificador calcula probabilidades para as classes. "
        "O comite utiliza a media dessas probabilidades e escolhe a classe com maior valor final.",
    )
    add_paragraph(
        doc,
        "A votacao suave foi escolhida porque considera o grau de confianca de cada classificador, "
        "em vez de utilizar apenas a classe prevista. A combinacao busca equilibrar modelos com perfis "
        "diferentes de precisao e recall.",
    )

    doc.add_page_break()
    doc.add_heading("7. Resultados e metricas", level=1)
    add_table(
        doc,
        ["Modelo", "Acuracia", "Precisao", "Recall", "F1-score"],
        [
            ("Comite", "0,7743", "0,5603", "0,6952", "0,6205"),
            ("Naive Bayes", "0,6948", "0,4589", "0,8369", "0,5928"),
            ("KNN", "0,7779", "0,5822", "0,5775", "0,5799"),
            ("Arvore de Decisao", "0,7970", "0,6679", "0,4679", "0,5503"),
            ("SVM", "0,7935", "0,6627", "0,4519", "0,5374"),
        ],
        [1.9, 1.15, 1.15, 1.15, 1.15],
    )
    add_caption(doc, "Tabela 1 - Comparacao dos modelos no conjunto de teste. Fonte: elaboracao propria.")
    doc.add_picture(str(metrics_path), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figura 2 - Comparacao das metricas. Fonte: elaboracao propria.")

    doc.add_heading("8. Discussao", level=1)
    add_paragraph(
        doc,
        "A Arvore de Decisao apresentou a maior acuracia individual (0,7970), mas identificou apenas "
        "46,79% dos clientes que realmente cancelaram. Esse resultado mostra por que a acuracia isolada "
        "pode ser enganosa em bases desbalanceadas.",
    )
    add_paragraph(
        doc,
        "O Naive Bayes apresentou o maior recall (0,8369), identificando a maior parte dos clientes com "
        "churn. Entretanto, sua precisao de 0,4589 indica maior quantidade de falsos alertas. O comite "
        "reduziu esse desequilibrio: atingiu recall de 0,6952, precisao de 0,5603 e o melhor F1-score, "
        "0,6205. Assim, o comite apresentou melhora sobre o melhor classificador individual segundo o "
        "criterio principal adotado.",
    )
    doc.add_heading("8.1 Limitacoes", level=2)
    add_bullets(
        doc,
        [
            "Os resultados representam uma unica divisao entre treino e teste.",
            "Os hiperparametros foram definidos de forma inicial, sem uma busca extensa.",
            "A base nao informa o momento exato do cancelamento nem os custos das acoes de retencao.",
            "O desbalanceamento influencia o comportamento dos classificadores.",
        ],
    )
    doc.add_heading("8.2 Melhorias futuras", level=2)
    add_bullets(
        doc,
        [
            "Aplicar validacao cruzada e busca de hiperparametros.",
            "Avaliar pesos diferentes para cada integrante do comite.",
            "Testar tecnicas de balanceamento e ajuste do limiar de classificacao.",
            "Incluir metricas como curva ROC, AUC e curva precisao-recall.",
            "Investigar interpretabilidade e importancia das features.",
        ],
    )

    doc.add_heading("9. Conclusao", level=1)
    add_paragraph(
        doc,
        "O projeto cumpriu o objetivo de construir um pipeline completo de classificacao, desde a "
        "analise e preparacao dos dados ate a avaliacao de modelos individuais e de um comite. Os "
        "resultados confirmaram que modelos com maior acuracia nem sempre sao os mais adequados para "
        "identificar clientes em risco de cancelamento.",
    )
    add_paragraph(
        doc,
        "O comite por votacao suave apresentou o melhor F1-score e um equilibrio superior entre "
        "precisao e recall. Portanto, para o problema proposto, a combinacao dos classificadores foi "
        "mais adequada do que a utilizacao isolada dos modelos avaliados.",
    )

    doc.add_heading("Referencias", level=1)
    add_paragraph(
        doc,
        "KAGGLE. Telco Customer Churn. Disponivel em: https://www.kaggle.com/datasets/blastchar/"
        "telco-customer-churn. Acesso em: 9 jun. 2026.",
    )
    add_paragraph(
        doc,
        "SCIKIT-LEARN DEVELOPERS. Scikit-learn: machine learning in Python. Disponivel em: "
        "https://scikit-learn.org/. Acesso em: 9 jun. 2026.",
    )

    output = REPORTS / "Relatorio_Projeto_A3_Comite_Classificadores.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    print(build_report())
