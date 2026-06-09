from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "Roteiro_Explicacao_Projeto_A3.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 42, 55)
GRAY = RGBColor(90, 98, 108)


def font(run, size=11, bold=False, italic=False, color=DARK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for margin, value in (("top", 100), ("start", 130), ("bottom", 100), ("end", 130)):
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Projeto A3 | ")
    font(run, size=9, color=GRAY)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    font(run, size=9, color=GRAY)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def para(doc, text, lead=None, italic=False):
    p = doc.add_paragraph()
    if lead and text.startswith(lead):
        r = p.add_run(lead)
        font(r, bold=True)
        r = p.add_run(text[len(lead) :])
        font(r, italic=italic)
    else:
        r = p.add_run(text)
        font(r, italic=italic)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        font(r)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(item)
        font(r)


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, text in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade(cell, "E8EEF5")
        cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        font(r, size=9.5, bold=True, color=BLUE)
    for values in rows:
        row = tbl.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            font(r, size=9.5, bold=(i == 0))
    return tbl


def build():
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(15)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("ROTEIRO DE EXPLICACAO DO PROJETO")
    font(r, size=22, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    r = subtitle.add_run("Comite de classificadores para previsao de churn")
    font(r, size=13, color=GRAY)

    doc.add_heading("1. Qual problema meu projeto resolve?", level=1)
    para(
        doc,
        "Meu projeto tenta prever quais clientes de uma empresa de telecomunicacoes possuem maior "
        "risco de cancelar o servico. Esse cancelamento e chamado de churn. Identificar esses clientes "
        "antecipadamente pode ajudar a empresa a criar acoes de retencao, oferecer melhores condicoes "
        "e evitar perdas de receita.",
    )
    para(
        doc,
        "A tarefa e um problema de classificacao binaria: o cliente pertence a classe Churn = Yes "
        "quando cancelou e Churn = No quando permaneceu.",
    )

    doc.add_heading("2. O que meu projeto faz?", level=1)
    para(
        doc,
        "O projeto analisa os dados dos clientes, prepara esses dados para uso por algoritmos de "
        "inteligencia artificial, treina quatro classificadores diferentes e depois combina os modelos "
        "em um comite. Por fim, compara o desempenho dos modelos para descobrir qual abordagem apresenta "
        "o melhor resultado.",
    )
    bullets(
        doc,
        [
            "Carrega e analisa uma base publica do Kaggle com 7.043 clientes.",
            "Trata valores ausentes e transforma variaveis categoricas em numeros.",
            "Separa os dados entre treino e teste.",
            "Treina KNN, Naive Bayes, SVM e Arvore de Decisao.",
            "Combina os quatro modelos em um comite por votacao suave.",
            "Avalia acuracia, precisao, recall, F1-score e matrizes de confusao.",
        ],
    )

    doc.add_heading("3. Como eu fiz o projeto", level=1)
    numbered(
        doc,
        [
            "Escolhi a base Telco Customer Churn no Kaggle porque ela possui uma variavel alvo clara e representa um problema real de classificacao.",
            "Defini Churn como target e removi customerID, pois o identificador nao ajuda a prever cancelamento.",
            "Analisei a distribuicao do target e observei que a base e desbalanceada: existem mais clientes que nao cancelaram.",
            "Converti TotalCharges para numero e tratei seus 11 valores ausentes usando a mediana.",
            "Codifiquei as variaveis categoricas com one-hot encoding e padronizei as variaveis numericas.",
            "Dividi os dados de forma estratificada: 80% para treino e 20% para teste.",
            "Treinei e avaliei os quatro classificadores no mesmo conjunto de teste.",
            "Criei um comite por votacao suave, que calcula a media das probabilidades previstas pelos modelos.",
        ],
    )

    doc.add_heading("4. Por que usei um pipeline?", level=1)
    para(
        doc,
        "Utilizei pipelines para unir o pre-processamento e o classificador. Isso garante que as mesmas "
        "transformacoes sejam aplicadas corretamente aos dados e evita vazamento de dados, pois o "
        "pre-processamento e ajustado somente com o conjunto de treino.",
    )

    doc.add_heading("5. Classificadores utilizados", level=1)
    table(
        doc,
        ["Modelo", "Como funciona de forma resumida"],
        [
            ("KNN", "Classifica usando os clientes mais parecidos e proximos."),
            ("Naive Bayes", "Calcula probabilidades com base nas caracteristicas observadas."),
            ("SVM", "Procura uma fronteira que separe as classes."),
            ("Arvore de Decisao", "Cria regras e decisoes sucessivas sobre as features."),
            ("Comite", "Combina as probabilidades dos quatro modelos."),
        ],
        [1.65, 4.75],
    )

    doc.add_heading("6. Como avaliar os resultados", level=1)
    bullets(
        doc,
        [
            "Acuracia: percentual total de previsoes corretas.",
            "Precisao: entre os clientes previstos como churn, quantos realmente cancelaram.",
            "Recall: entre os clientes que cancelaram, quantos foram identificados.",
            "F1-score: equilibrio entre precisao e recall.",
            "Matriz de confusao: mostra acertos e erros para cada classe.",
        ],
    )
    para(
        doc,
        "Como a base e desbalanceada, escolhi o F1-score como principal criterio de comparacao. Uma "
        "acuracia alta pode esconder o fato de que o modelo nao identifica os clientes que realmente cancelam.",
    )

    doc.add_heading("7. Resultados obtidos", level=1)
    table(
        doc,
        ["Modelo", "Acuracia", "Precisao", "Recall", "F1-score"],
        [
            ("Comite", "0,7743", "0,5603", "0,6952", "0,6205"),
            ("Naive Bayes", "0,6948", "0,4589", "0,8369", "0,5928"),
            ("KNN", "0,7779", "0,5822", "0,5775", "0,5799"),
            ("Arvore", "0,7970", "0,6679", "0,4679", "0,5503"),
            ("SVM", "0,7935", "0,6627", "0,4519", "0,5374"),
        ],
        [1.55, 1.2, 1.2, 1.2, 1.2],
    )
    para(
        doc,
        "O melhor modelo individual pelo F1-score foi o Naive Bayes, com 0,5928. O comite obteve "
        "F1-score de 0,6205 e apresentou melhora sobre todos os modelos individuais. A Arvore teve a "
        "maior acuracia, mas um recall baixo, mostrando que acuracia isolada nao e suficiente.",
    )

    doc.add_heading("8. O que mostrar ao professor", level=1)
    bullets(
        doc,
        [
            "Abrir o notebook 01_eda_preprocessamento.ipynb e mostrar o carregamento da base, os graficos e o tratamento dos dados.",
            "Mostrar que o pre-processamento terminou com 45 atributos e zero valores ausentes.",
            "Abrir o notebook 02_modelagem_comite.ipynb e mostrar a definicao dos quatro modelos.",
            "Executar o notebook e mostrar a tabela de metricas e as matrizes de confusao.",
            "Finalizar mostrando que o comite atingiu F1-score 0,6205, superior ao melhor modelo individual.",
        ],
    )

    doc.add_heading("9. Fala curta para a demonstracao", level=1)
    para(
        doc,
        '"O problema escolhido foi prever o churn de clientes de telecomunicacoes. Primeiro, analisei '
        "e preparei a base, tratando valores ausentes, codificando categorias e padronizando valores "
        "numericos. Depois, treinei quatro classificadores diferentes e criei um comite por votacao "
        "suave. Como a base e desbalanceada, utilizei principalmente o F1-score para comparar os modelos. "
        "O melhor modelo individual foi o Naive Bayes, com F1-score de 0,5928, enquanto o comite atingiu "
        '0,6205. Portanto, a combinacao dos classificadores apresentou o melhor equilibrio geral."',
        italic=True,
    )

    doc.add_heading("10. Limitacoes e proximos passos", level=1)
    bullets(
        doc,
        [
            "Os modelos foram avaliados em uma unica divisao de treino e teste.",
            "Os hiperparametros podem ser melhorados com busca automatizada.",
            "O comite pode testar pesos diferentes para cada classificador.",
            "Podem ser avaliadas tecnicas de balanceamento e interpretabilidade.",
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
